"""
Document Agent tools: text extraction, summarization, entity extraction.

Libraries:
  pypdf       — PDF text extraction (pure Python, no system deps)
  python-docx — DOCX extraction
  httpx       — URL fetching
  BeautifulSoup — HTML cleaning

Design notes:
  - All public functions are async (called from execute_mcp_tool)
  - Heavy I/O (file reads, HTTP) uses asyncio.to_thread or httpx async
  - summarize_document uses Gemini directly — no MCP round-trip needed
    because this tool IS the LLM call (not a tool the LLM calls)
  - extract_entities uses regex + heuristics (no spaCy dep = no model download)
    Good enough for portfolio; note limitation in docstring
"""

from __future__ import annotations

import asyncio
import io
import json
import re
import tempfile
from dataclasses import dataclass, field
from pathlib import Path
from typing import Any
from urllib.parse import urlparse

import httpx
import structlog
from bs4 import BeautifulSoup

log = structlog.get_logger(__name__)


# ---------------------------------------------------------------------------
# Result dataclasses
# ---------------------------------------------------------------------------

@dataclass
class ExtractionResult:
    success: bool
    text: str
    source: str                      # "pdf" | "docx" | "url" | "txt"
    page_count: int = 0
    word_count: int = 0
    char_count: int = 0
    error: str = ""

    def summary_line(self) -> str:
        if not self.success:
            return f"Extraction failed: {self.error}"
        return (
            f"Source: {self.source} | "
            f"Pages: {self.page_count} | "
            f"Words: {self.word_count} | "
            f"Chars: {self.char_count}"
        )


@dataclass
class EntityResult:
    success: bool
    entities: dict[str, list[str]] = field(default_factory=dict)
    # keys: "names", "emails", "urls", "dates", "organizations", "numbers"
    total_found: int = 0
    error: str = ""


@dataclass
class SummaryResult:
    success: bool
    summary: str = ""
    key_points: list[str] = field(default_factory=list)
    word_count_original: int = 0
    word_count_summary: int = 0
    error: str = ""


# ---------------------------------------------------------------------------
# Internal helpers
# ---------------------------------------------------------------------------

def _word_count(text: str) -> int:
    return len(text.split())


def _truncate_for_llm(text: str, max_chars: int = 12_000) -> str:
    """
    Truncate text to fit Gemini context window safely.

    12k chars ≈ 3k tokens — well within flash-lite limits.
    Truncate at sentence boundary if possible.
    """
    if len(text) <= max_chars:
        return text
    truncated = text[:max_chars]
    # Try to end at last sentence boundary
    last_period = truncated.rfind(".")
    if last_period > max_chars * 0.8:
        truncated = truncated[: last_period + 1]
    return truncated + "\n\n[... text truncated for processing ...]"


def _extract_pdf_sync(data: bytes) -> tuple[str, int]:
    """
    Sync PDF extraction — runs in thread pool.
    Returns (text, page_count).
    Raises on corrupt/encrypted PDF.
    """
    import pypdf  # local import — only loaded when needed

    reader = pypdf.PdfReader(io.BytesIO(data))
    pages = []
    for page in reader.pages:
        page_text = page.extract_text() or ""
        pages.append(page_text)
    return "\n\n".join(pages), len(pages)


def _extract_docx_sync(data: bytes) -> tuple[str, int]:
    """
    Sync DOCX extraction — runs in thread pool.
    Returns (text, paragraph_count_as_page_proxy).
    """
    import docx  # python-docx

    doc = docx.Document(io.BytesIO(data))
    paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
    # Also grab table cells
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text.strip():
                    paragraphs.append(cell.text.strip())
    text = "\n\n".join(paragraphs)
    # No real "pages" in DOCX — approximate from paragraphs
    page_proxy = max(1, len(paragraphs) // 15)
    return text, page_proxy


def _clean_html(html: str, base_url: str = "") -> str:
    """
    Extract readable text from HTML.

    Strategy:
    1. Remove script/style/nav/footer/header tags entirely
    2. Prefer <article> or <main> if present (main content)
    3. Fall back to <body>
    4. get_text() with newline separator
    5. Collapse blank lines
    """
    soup = BeautifulSoup(html, "html.parser")

    # Remove noise tags
    for tag in soup(["script", "style", "nav", "footer", "header",
                     "aside", "form", "noscript", "iframe"]):
        tag.decompose()

    # Prefer semantic content containers
    content = (
        soup.find("article")
        or soup.find("main")
        or soup.find(id=re.compile(r"content|article|post|main", re.I))
        or soup.find(class_=re.compile(r"content|article|post|main", re.I))
        or soup.find("body")
        or soup
    )

    text = content.get_text(separator="\n")  # type: ignore[union-attr]

    # Collapse multiple blank lines → max 2
    text = re.sub(r"\n{3,}", "\n\n", text)
    # Strip leading whitespace per line
    lines = [line.strip() for line in text.splitlines()]
    text = "\n".join(line for line in lines if line)

    return text


# ---------------------------------------------------------------------------
# Public tool functions
# ---------------------------------------------------------------------------

async def extract_text(
    source: str,
    source_type: str = "auto",
) -> ExtractionResult:
    """
    Extract text from a PDF file path, DOCX file path, or URL.

    Args:
        source: File path (e.g. "/tmp/doc.pdf") or URL
                (e.g. "https://example.com/article")
        source_type: "auto" | "pdf" | "docx" | "url" | "txt"
                     "auto" detects from extension or URL scheme.

    Returns:
        ExtractionResult with .text and metadata.
    """
    # --- Detect type ---
    if source_type == "auto":
        parsed = urlparse(source)
        if parsed.scheme in ("http", "https"):
            source_type = "url"
        elif source.lower().endswith(".pdf"):
            source_type = "pdf"
        elif source.lower().endswith(".docx"):
            source_type = "docx"
        elif source.lower().endswith(".txt"):
            source_type = "txt"
        else:
            source_type = "url"  # default — treat unknown as URL

    log.info("extract_text", source=source[:80], source_type=source_type)

    try:
        # --- PDF ---
        if source_type == "pdf":
            path = Path(source)
            if not path.exists():
                return ExtractionResult(
                    success=False,
                    text="",
                    source="pdf",
                    error=f"File not found: {source}",
                )
            data = path.read_bytes()
            text, pages = await asyncio.to_thread(_extract_pdf_sync, data)
            return ExtractionResult(
                success=True,
                text=text,
                source="pdf",
                page_count=pages,
                word_count=_word_count(text),
                char_count=len(text),
            )

        # --- DOCX ---
        elif source_type == "docx":
            path = Path(source)
            if not path.exists():
                return ExtractionResult(
                    success=False,
                    text="",
                    source="docx",
                    error=f"File not found: {source}",
                )
            data = path.read_bytes()
            text, pages = await asyncio.to_thread(_extract_docx_sync, data)
            return ExtractionResult(
                success=True,
                text=text,
                source="docx",
                page_count=pages,
                word_count=_word_count(text),
                char_count=len(text),
            )

        # --- TXT ---
        elif source_type == "txt":
            path = Path(source)
            if not path.exists():
                return ExtractionResult(
                    success=False,
                    text="",
                    source="txt",
                    error=f"File not found: {source}",
                )
            text = await asyncio.to_thread(path.read_text, encoding="utf-8",
                                           errors="replace")
            return ExtractionResult(
                success=True,
                text=text,
                source="txt",
                page_count=1,
                word_count=_word_count(text),
                char_count=len(text),
            )

        # --- URL ---
        elif source_type == "url":
            async with httpx.AsyncClient(
                timeout=15.0,
                follow_redirects=True,
                headers={"User-Agent": "Mozilla/5.0 (A2A Document Agent)"},
            ) as client:
                response = await client.get(source)
                response.raise_for_status()

            content_type = response.headers.get("content-type", "")

            if "pdf" in content_type:
                # URL returned a PDF
                text, pages = await asyncio.to_thread(
                    _extract_pdf_sync, response.content
                )
                return ExtractionResult(
                    success=True,
                    text=text,
                    source="pdf",
                    page_count=pages,
                    word_count=_word_count(text),
                    char_count=len(text),
                )
            else:
                # Assume HTML
                text = _clean_html(response.text, base_url=source)
                return ExtractionResult(
                    success=True,
                    text=text,
                    source="url",
                    page_count=1,
                    word_count=_word_count(text),
                    char_count=len(text),
                )

        else:
            return ExtractionResult(
                success=False,
                text="",
                source=source_type,
                error=f"Unknown source_type: {source_type}",
            )

    except httpx.HTTPStatusError as e:
        log.warning("extract_text_http_error", status=e.response.status_code)
        return ExtractionResult(
            success=False,
            text="",
            source=source_type,
            error=f"HTTP {e.response.status_code}: {source}",
        )
    except Exception as e:
        log.exception("extract_text_error", source=source[:80])
        return ExtractionResult(
            success=False,
            text="",
            source=source_type,
            error=str(e),
        )


async def summarize_document(
    text: str,
    style: str = "concise",
    max_length: int = 500,
) -> SummaryResult:
    """
    Summarize document text using Gemini.

    Args:
        text: Document text to summarize.
        style: "concise" | "detailed" | "bullet" | "executive"
        max_length: Target word count for summary (approximate).

    Returns:
        SummaryResult with .summary and .key_points list.

    Note: This calls Gemini directly. When used as an MCP tool,
    the agent's LLM calls this via execute_mcp_tool → Gemini makes
    a nested LLM call. That's intentional: summarization IS the task,
    not a step toward another LLM response.
    """
    from google import genai
    from google.genai import types as genai_types
    from shared.config import settings

    if not text.strip():
        return SummaryResult(
            success=False,
            error="No text provided to summarize",
        )

    truncated = _truncate_for_llm(text, max_chars=12_000)
    word_count_orig = _word_count(text)

    style_instructions = {
        "concise": (
            f"Write a concise summary in approximately {max_length} words. "
            "Focus on the main points only."
        ),
        "detailed": (
            f"Write a detailed summary in approximately {max_length} words. "
            "Cover all major sections and supporting details."
        ),
        "bullet": (
            "Write a bullet-point summary. "
            "Start with a 1-sentence overview, then 5-10 key bullet points. "
            f"Total approximately {max_length} words."
        ),
        "executive": (
            f"Write an executive summary in approximately {max_length} words. "
            "Focus on: purpose, key findings, recommendations, and next steps."
        ),
    }.get(style, f"Summarize the following text in approximately {max_length} words.")

    prompt = (
        f"{style_instructions}\n\n"
        f"Also extract exactly 3-5 key points as a JSON list of strings "
        f"at the END of your response in this format:\n"
        f"KEY_POINTS_JSON: [\"point 1\", \"point 2\", \"point 3\"]\n\n"
        f"Document text:\n{truncated}"
    )

    log.info("summarize_document", style=style, words=word_count_orig)

    try:
        client = genai.Client(api_key=settings.google_api_key)

        response = await asyncio.to_thread(
            client.models.generate_content,
            model=settings.agent_model,
            contents=prompt,
            config=genai_types.GenerateContentConfig(
                temperature=0.3,
                system_instruction=(
                    "You are a document summarization expert. "
                    "Produce clear, accurate summaries. "
                    "Always include the KEY_POINTS_JSON line at the end."
                ),
            ),
        )

        full_text = response.text or ""

        # Parse key points from structured marker
        key_points: list[str] = []
        summary = full_text

        kp_match = re.search(
            r"KEY_POINTS_JSON:\s*(\[.*?\])",
            full_text,
            re.DOTALL,
        )
        if kp_match:
            try:
                key_points = json.loads(kp_match.group(1))
                # Remove the KEY_POINTS_JSON line from summary text
                summary = full_text[: kp_match.start()].strip()
            except json.JSONDecodeError:
                # Keep full text as summary, no key points
                summary = full_text

        return SummaryResult(
            success=True,
            summary=summary,
            key_points=key_points,
            word_count_original=word_count_orig,
            word_count_summary=_word_count(summary),
        )

    except Exception as e:
        log.exception("summarize_document_error")
        return SummaryResult(success=False, error=str(e))


async def extract_entities(
    text: str,
    entity_types: list[str] | None = None,
) -> EntityResult:
    """
    Extract named entities from text using regex heuristics.

    Entity types supported:
      "emails"    — email addresses
      "urls"      — http/https URLs
      "dates"     — common date patterns (2024-01-15, Jan 15 2024, etc.)
      "numbers"   — dollar amounts, percentages, large numbers
      "names"     — capitalized word sequences (heuristic, not NLP)
      "organizations" — words before Inc/Corp/LLC/Ltd/Co (heuristic)

    Note: This uses regex, not spaCy/NER. Avoids model download requirement.
    Precision over recall — false positives minimized, may miss some entities.

    Args:
        text: Document text to analyze.
        entity_types: List of entity types to extract, or None for all.

    Returns:
        EntityResult with .entities dict keyed by type.
    """
    if not text.strip():
        return EntityResult(success=False, error="No text provided")

    all_types = ["emails", "urls", "dates", "numbers", "names", "organizations"]
    types_to_extract = entity_types or all_types

    # Validate requested types
    invalid = set(types_to_extract) - set(all_types)
    if invalid:
        return EntityResult(
            success=False,
            error=f"Unknown entity types: {invalid}. Valid: {all_types}",
        )

    log.info("extract_entities", types=types_to_extract, chars=len(text))

    entities: dict[str, list[str]] = {}

    # --- Emails ---
    if "emails" in types_to_extract:
        pattern = r"\b[A-Za-z0-9._%+\-]+@[A-Za-z0-9.\-]+\.[A-Za-z]{2,}\b"
        found = list(dict.fromkeys(re.findall(pattern, text)))  # unique, order-preserving
        entities["emails"] = found

    # --- URLs ---
    if "urls" in types_to_extract:
        pattern = r"https?://[^\s\)\]\"\'<>]{5,}"
        found = list(dict.fromkeys(re.findall(pattern, text)))
        # Clean trailing punctuation
        found = [u.rstrip(".,;:)]}") for u in found]
        entities["urls"] = found

    # --- Dates ---
    if "dates" in types_to_extract:
        date_patterns = [
            # ISO: 2024-01-15
            r"\b\d{4}-\d{2}-\d{2}\b",
            # US: 01/15/2024 or 1/15/24
            r"\b\d{1,2}/\d{1,2}/\d{2,4}\b",
            # Written: January 15, 2024 or Jan 15, 2024
            r"\b(?:January|February|March|April|May|June|July|August|"
            r"September|October|November|December|"
            r"Jan|Feb|Mar|Apr|Jun|Jul|Aug|Sep|Oct|Nov|Dec)"
            r"\s+\d{1,2},?\s+\d{4}\b",
            # Written: 15 January 2024
            r"\b\d{1,2}\s+(?:January|February|March|April|May|June|July|"
            r"August|September|October|November|December)\s+\d{4}\b",
            # Quarter: Q1 2024
            r"\bQ[1-4]\s+\d{4}\b",
        ]
        found_dates: list[str] = []
        for pat in date_patterns:
            found_dates.extend(re.findall(pat, text, re.IGNORECASE))
        entities["dates"] = list(dict.fromkeys(found_dates))

    # --- Numbers (monetary, percentages, significant figures) ---
    if "numbers" in types_to_extract:
        num_patterns = [
            # Dollar amounts: $1,234.56 or $1.2M or $500K
            r"\$[\d,]+(?:\.\d+)?(?:[KMBTkmbt](?:illion|illions?)?)?\b",
            # Percentages: 15.3% or 15%
            r"\b\d+(?:\.\d+)?%",
            # Large numbers with commas: 1,234,567
            r"\b\d{1,3}(?:,\d{3}){2,}\b",
        ]
        found_nums: list[str] = []
        for pat in num_patterns:
            found_nums.extend(re.findall(pat, text))
        entities["numbers"] = list(dict.fromkeys(found_nums))

    # --- Names (heuristic: 2-4 consecutive capitalized words) ---
    if "names" in types_to_extract:
        # Match sequences of 2-4 Title Case words
        # Exclude common false positives (sentence starts caught by context)
        pattern = r"\b([A-Z][a-z]+(?:\s+[A-Z][a-z]+){1,3})\b"
        candidates = re.findall(pattern, text)

        # Filter out common non-name phrases
        stop_phrases = {
            "The", "This", "That", "These", "Those", "In", "On", "At",
            "For", "With", "From", "By", "As", "An", "New", "United States",
            "United Kingdom", "European Union",
        }
        names = [
            c for c in candidates
            if c not in stop_phrases
            and not any(c.startswith(s) for s in ["The ", "This ", "That "])
            and len(c.split()) >= 2  # at least 2 words
        ]
        entities["names"] = list(dict.fromkeys(names))[:50]  # cap at 50

    # --- Organizations (heuristic: words before Inc/Corp/LLC etc.) ---
    if "organizations" in types_to_extract:
        pattern = (
            r"\b([A-Z][A-Za-z\s&,\.]{2,40}?)"
            r"(?:\s+(?:Inc\.?|Corp\.?|LLC|Ltd\.?|Co\.?|"
            r"Company|Corporation|Limited|Group|"
            r"Foundation|Institute|Association|University|College))"
        )
        found_orgs = re.findall(pattern, text)
        # Also catch ALL-CAPS acronyms (IBM, NASA, etc.)
        acronyms = re.findall(r"\b[A-Z]{2,6}\b", text)
        # Filter common word acronyms
        word_acronyms = {"US", "UK", "EU", "UN", "GDP", "CEO", "CFO",
                         "CTO", "API", "URL", "PDF", "AI", "ML", "IT"}
        real_acronyms = [a for a in acronyms if a not in word_acronyms]

        all_orgs = list(dict.fromkeys(
            [o.strip() for o in found_orgs] + real_acronyms
        ))
        entities["organizations"] = all_orgs[:30]  # cap at 30

    total = sum(len(v) for v in entities.values())
    return EntityResult(
        success=True,
        entities=entities,
        total_found=total,
    )